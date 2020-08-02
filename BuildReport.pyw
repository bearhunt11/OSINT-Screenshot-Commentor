import os
import BasicInfo
from docx import Document
from docx.shared import Inches
import fnmatch

# file contains the filename of the .bh11 comment file
def create_report(file, folder):
    print('Report builder started:')
    list_folder = fnmatch.filter(os.listdir(folder), "*.bh11")
    print(list_folder)

    # create WORD Document and write content, create numbered screenshots
    document = Document()
    document.add_heading('Screenshot & Comments', 0)
    scr_nr = 1
    for file in list_folder:
        comment_file = file
        scr_file = file.replace('.bh11', '.png')
        full_comment_file = os.path.join(folder, comment_file)
        full_scr_file = os.path.join(folder, scr_file)

        with open(full_comment_file, 'r') as f:
            full_comment_txt = f.readlines()

        osint_header = ""
        for line in full_comment_txt[1:5]:
            osint_header += line

        osint_comment = ""
        for line in full_comment_txt[8:]:
            osint_comment += line


        ### Create WORD content
        hdr = 'Screenshot ' + str(scr_nr) + ':'
        document.add_heading(hdr, level=1)
        document.add_picture(full_scr_file, width=Inches(4.50))
        document.add_heading('Basic screenshot information:', level=2)
        document.add_paragraph(osint_header)
        document.add_heading('OSINT Comment:', level=2)
        document.add_paragraph(osint_comment)
        document.add_page_break()

        # Add 1 to the screenshot counter
        scr_nr += 1

    # Close and save WORD file
    report = os.path.join(folder, 'ScreenshotReport.docx')
    document.save(report)
    print('Report finished')
